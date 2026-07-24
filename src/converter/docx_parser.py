import os
from typing import List
from .base import BaseParser, DocumentModel, DocumentMetadata, Section, TableData

class DOCXParser(BaseParser):
    SUPPORTED_EXTENSIONS = {'.docx', '.doc'}

    def can_parse(self, file_extension: str) -> bool:
        return file_extension.lower() in self.SUPPORTED_EXTENSIONS

    def parse(self, file_path: str) -> DocumentModel:
        file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0
        filename = os.path.basename(file_path)

        sections: List[Section] = []
        tables: List[TableData] = []
        unstructured_elements = []
        raw_text_parts = []

        try:
            import docx
            doc = docx.Document(file_path)

            current_section = Section(title="Document Body", level=1, content="")
            curr_lines = []

            for paragraph in doc.paragraphs:
                p_text = paragraph.text.strip()
                if not p_text:
                    continue

                raw_text_parts.append(p_text)
                style_name = paragraph.style.name.lower() if paragraph.style else ""

                if 'heading' in style_name:
                    if curr_lines:
                        current_section.content = "\n".join(curr_lines)
                        sections.append(current_section)
                        curr_lines = []

                    level = 1
                    try:
                        level_str = ''.join(c for c in style_name if c.isdigit())
                        if level_str:
                            level = int(level_str)
                    except ValueError:
                        level = 1

                    current_section = Section(title=p_text, level=level, content="")
                    unstructured_elements.append({"type": "heading", "level": level, "text": p_text})
                elif 'list' in style_name or 'bullet' in style_name:
                    curr_lines.append(f"- {p_text}")
                    unstructured_elements.append({"type": "list_item", "text": p_text})
                else:
                    curr_lines.append(p_text)
                    unstructured_elements.append({"type": "paragraph", "text": p_text})

            if curr_lines:
                current_section.content = "\n".join(curr_lines)
                sections.append(current_section)

            # Parse docx tables
            for idx, table in enumerate(doc.tables):
                headers = []
                rows = []
                for r_idx, row in enumerate(table.rows):
                    row_cells = [cell.text.strip() for cell in row.cells]
                    if r_idx == 0:
                        headers = row_cells
                    else:
                        rows.append(row_cells)
                if headers or rows:
                    t_obj = TableData(headers=headers, rows=rows, caption=f"Table {idx+1}")
                    tables.append(t_obj)

        except Exception as e:
            sections.append(Section(
                title="Error Parsing DOCX",
                level=1,
                content=f"Could not parse docx file: {str(e)}"
            ))

        full_raw_text = "\n\n".join(raw_text_parts)
        char_count = len(full_raw_text)
        word_count = len(full_raw_text.split())
        line_count = len(full_raw_text.splitlines())

        metadata = DocumentMetadata(
            filename=filename,
            file_type="DOCX",
            file_size_bytes=file_size,
            char_count=char_count,
            word_count=word_count,
            line_count=line_count
        )

        doc_title = os.path.splitext(filename)[0].replace('_', ' ').replace('-', ' ').title()

        return DocumentModel(
            metadata=metadata,
            title=doc_title,
            sections=sections,
            tables=tables,
            raw_text=full_raw_text,
            unstructured_elements=unstructured_elements
        )
